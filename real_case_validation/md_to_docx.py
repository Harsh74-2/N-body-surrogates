"""
md_to_docx.py
=============
Combine every audit Markdown report in `real_case_validation/` into a
single Word document, with a TOC table at the top. Handles headings,
bullets, numbered lists, bold/inline-code, code blocks, AND tables
(markdown pipe tables are the bulk of the audit content).

The output is meant to be ready to paste into a report appendix —
tables retain their alignment, code blocks use a monospace font, and
the TOC at the top lets a reader jump straight to any section.

Usage
-----
    python -m real_case_validation.md_to_docx \\
        --out  real_case_validation/report_bundle.docx

Reads, in order:
  1. cross_N_audit.md            (autoregressive rollout cross-N)
  2. cross_N_audit_single_step.md (single-step bare-error cross-N)
  3. report_N10/N10_audit.md     (per-N standalone)
  4. report_N25/N25_audit.md
  5. report_N50/N50_audit.md
  6. report_N100/N100_audit.md
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ── Reading order ───────────────────────────────────────────────────────
DEFAULT_ORDER = [
    "cross_N_audit.md",
    "cross_N_audit_single_step.md",
    "report_N10/N10_audit.md",
    "report_N25/N25_audit.md",
    "report_N50/N50_audit.md",
    "report_N100/N100_audit.md",
]


def _md_to_blocks(md_text: str) -> list[dict]:
    """
    Lightweight Markdown -> block stream. Each block is one of:
      {"type": "h1"|"h2"|"h3", "text": "..."}
      {"type": "p",            "text": "..."}
      {"type": "ul"|"ol",     "items": ["...", "..."]}
      {"type": "code",         "text": "..."}
      {"type": "table",        "rows": [["..."], ...], "header": [...]}
    Tables are detected by scanning a contiguous run of `|` lines.
    """
    lines = md_text.splitlines()
    blocks: list[dict] = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── Blank ──────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Heading ────────────────────────────────────────────
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:].strip()})
            i += 1; continue
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:].strip()})
            i += 1; continue
        if stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:].strip()})
            i += 1; continue

        # ── Code block ────────────────────────────────────────
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].lstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1  # skip closing fence
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        # ── Table ─────────────────────────────────────────────
        # A markdown pipe table has at least a header row, a
        # separator row (---), and at least one body row.
        if "|" in stripped and i + 1 < n:
            # Peek ahead: is the next line a separator?
            next_line = lines[i + 1].strip()
            if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", next_line):
                # Found a table; consume all consecutive pipe rows.
                tbl_lines = []
                while i < n and "|" in lines[i] and lines[i].strip():
                    tbl_lines.append(lines[i])
                    i += 1
                # Parse: row 0 = header, row 1 = separator, rest = body.
                rows = []
                for tl in tbl_lines:
                    cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                    rows.append(cells)
                if len(rows) >= 2:
                    header = rows[0]
                    body   = rows[2:]
                    blocks.append({"type": "table",
                                   "header": header, "rows": body})
                continue

        # ── Bulleted list ─────────────────────────────────────
        if stripped.startswith("- "):
            items = []
            while i < n and lines[i].lstrip().startswith("- "):
                items.append(lines[i].lstrip()[2:].strip())
                i += 1
            blocks.append({"type": "ul", "items": items})
            continue

        # ── Numbered list ─────────────────────────────────────
        if re.match(r"^\d+\.\s", stripped):
            items = []
            while i < n and re.match(r"^\s*\d+\.\s", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "ol", "items": items})
            continue

        # ── Plain paragraph (possibly multi-line) ─────────────
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and \
                not lines[i].lstrip().startswith(("#", "-", "```")) and \
                "|" not in lines[i]:
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "p", "text": "\n".join(para_lines)})

    return blocks


# ── Inline formatting helpers ──────────────────────────────────────────
def _add_runs(paragraph, text: str, *, bold: bool = False,
              mono: bool = False, size: int = 11) -> None:
    """Add runs to a paragraph, splitting on `**bold**` and ``code``.

    The `bold` and `mono` flags set the run style for the *whole*
    paragraph (used for headings). When the markdown source contains
    inline `**...**` or `` `...` `` markers, those parts inherit the
    same `bold` / `mono` style; the markers themselves are stripped
    (the [2:-2] / [1:-1] slice) and only the inner text becomes the
    run text.
    """
    # Tokenise on ** (bold) and ` (mono).
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        # Decide the run style. We track whether the run came from
        # a markdown **..** / `..` marker separately from the
        # caller's "this whole paragraph is bold" flag, because the
        # marker run has its delimiters (the `**` / `` ` ``) that
        # must be stripped, while the caller's flag does not.
        is_marker_bold = (part.startswith("**") and part.endswith("**")
                          and len(part) >= 4)
        is_marker_mono = (part.startswith("`") and part.endswith("`")
                          and len(part) >= 2)
        run_bold = bold or is_marker_bold
        run_mono = mono or is_marker_mono
        # Strip the markers if this part came from a marked-up
        # fragment; otherwise keep the entire part as-is.
        if is_marker_bold:
            run_text = part[2:-2]
        elif is_marker_mono:
            run_text = part[1:-1]
        else:
            run_text = part
        r = paragraph.add_run(run_text)
        r.bold = run_bold
        if run_mono:
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            # Ensure east-asian font fallback doesn't override.
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
        else:
            r.font.size = Pt(size)


# ── Table rendering ───────────────────────────────────────────────────
def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    """Add a markdown table as a styled Word table."""
    if not header:
        return
    n_cols = len(header)
    # Normalise body rows to the header width.
    body = []
    for r in rows:
        r2 = list(r) + [""] * (n_cols - len(r)) if len(r) < n_cols else list(r[:n_cols])
        body.append(r2)
    table = doc.add_table(rows=1 + len(body), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row.
    hdr = table.rows[0]
    for j, txt in enumerate(header):
        cell = hdr.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_runs(p, txt, bold=True, size=10)
        _shade_cell(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body rows.
    for ri, row in enumerate(body):
        tr = table.rows[ri + 1]
        for j, txt in enumerate(row):
            cell = tr.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, txt, size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ── Top-level builder ─────────────────────────────────────────────────
def build_docx(reports_root: Path, out_path: Path,
               order: list[str] | None = None,
               title: str = "Real-Case Validation Reports") -> None:
    """
    Combine the listed Markdown reports into a single Word document.

    The first page is a TOC table listing each report (with its source
    path); each report's content is appended under an H1 heading and
    page break.
    """
    if order is None:
        order = DEFAULT_ORDER

    doc = Document()

    # ── Page geometry: A4 with sensible margins ──────────────────────
    section = doc.sections[0]
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Title block ──────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(title_p, title, bold=True, size=20)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(sub_p,
              "Cross-N and per-N error-percentage audits for the "
              "real-life Solar-System surrogate validation. "
              "All six variants (MLP / LSTM / GNN × single + stable) "
              "are reported per preset.",
              size=11)
    doc.add_paragraph()

    # ── Resolve each report's path ─────────────────────────────────
    found: list[tuple[str, Path]] = []
    for rel in order:
        full = reports_root / rel
        if not full.exists():
            print(f"[warn] missing: {full}")
            continue
        found.append((rel, full))

    # ── TOC table ──────────────────────────────────────────────────
    toc_h = doc.add_paragraph()
    _add_runs(toc_h, "Contents", bold=True, size=14)
    toc_intro = doc.add_paragraph()
    _add_runs(toc_intro,
              "Every audit listed below is rendered in full in this "
              "document. Cross-N reports come first (they frame the "
              "side-by-side comparison); the four per-N audits come "
              "after (each is the standalone report for one training "
              "budget).",
              size=11)
    doc.add_paragraph()
    toc_rows = []
    for rel, full in found:
        # Short label from the path (basename without .md).
        label = full.stem
        if label == "cross_N_audit":
            label = "Cross-N Audit (autoregressive rollout)"
        elif label == "cross_N_audit_single_step":
            label = "Cross-N Audit (single-step bare error)"
        elif label.startswith("N") and label.endswith("_audit"):
            label = f"Per-N Audit (N = {label[1:-len('_audit')]})"
        toc_rows.append([label, str(rel)])
    _add_table(doc,
               header=["Report", "Source path"],
               rows=toc_rows)
    doc.add_page_break()

    # ── Body: each report in turn ─────────────────────────────────
    for rel, full in found:
        md = full.read_text(encoding="utf-8")
        blocks = _md_to_blocks(md)
        # Derive a short subtitle from the relative path so the
        # reader can tell two reports with the same H1 apart (the
        # cross-N rollout and single-step audits both use the same
        # first heading). The subtitle goes immediately after the
        # first H1 / paragraph block.
        subtitle = ""
        if rel == "cross_N_audit.md":
            subtitle = "(autoregressive rollout, mean err % per (N, model))"
        elif rel == "cross_N_audit_single_step.md":
            subtitle = "(single-step bare error, mean err % per (N, model))"
        elif rel.startswith("report_N") and "/N" in rel:
            # e.g. report_N25/N25_audit.md -> "N = 25"
            try:
                n_part = rel.split("report_N")[1].split("/")[0]
                subtitle = f"(per-N standalone audit, training budget N = {n_part})"
            except (IndexError, ValueError):
                pass
        subtitle_emitted = False
        for blk in blocks:
            t = blk["type"]
            if t == "h1":
                p = doc.add_paragraph()
                _add_runs(p, blk["text"], bold=True, size=18)
                if subtitle:
                    sp = doc.add_paragraph()
                    _add_runs(sp, subtitle, size=11)
                    subtitle_emitted = True
            elif t == "h2":
                p = doc.add_paragraph()
                _add_runs(p, blk["text"], bold=True, size=14)
            elif t == "h3":
                p = doc.add_paragraph()
                _add_runs(p, blk["text"], bold=True, size=12)
            elif t == "p":
                p = doc.add_paragraph()
                _add_runs(p, blk["text"], size=11)
            elif t == "ul":
                for item in blk["items"]:
                    p = doc.add_paragraph(style="List Bullet")
                    _add_runs(p, item, size=11)
            elif t == "ol":
                for item in blk["items"]:
                    p = doc.add_paragraph(style="List Number")
                    _add_runs(p, item, size=11)
            elif t == "code":
                p = doc.add_paragraph()
                _add_runs(p, blk["text"], mono=True, size=10)
            elif t == "table":
                _add_table(doc, header=blk["header"], rows=blk["rows"])
                doc.add_paragraph()  # breathing room after each table
        doc.add_page_break()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[docx] {out_path}  ({len(found)} reports, "
          f"{out_path.stat().st_size / 1024:.1f} kB)")


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--reports-root",
                    default="real_case_validation",
                    help="Directory containing the cross_N_audit*.md "
                         "and report_N*/N*_audit.md files.")
    ap.add_argument("--out", required=True,
                    help="Path to write the Word document.")
    ap.add_argument("--title", default="Real-Case Validation Reports",
                    help="Document title (top of page 1).")
    args = ap.parse_args()

    reports_root = Path(args.reports_root).expanduser().resolve()
    if not reports_root.exists():
        raise SystemExit(f"reports root not found: {reports_root}")
    out_path = Path(args.out).expanduser().resolve()
    build_docx(reports_root, out_path, title=args.title)


if __name__ == "__main__":
    main()